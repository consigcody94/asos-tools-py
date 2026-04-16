"""DOCX incident investigation report builder.

Given one or more ASOS stations and a time window, produces a
NOAA-styled Word document documenting every ``$`` maintenance-flag
incident in the window, per-station sensor code breakdowns, a root-cause
analysis of downstream alerting gaps, and concrete recommendations.

This module is intentionally self-contained so it can be called either
from a CLI (``deploy/generate_incident_report.py``) or from the Streamlit
app at runtime (buffered to ``io.BytesIO`` for browser download).
"""

from __future__ import annotations

import io
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Iterable, Union

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Pt, RGBColor

from asos_tools.metars import decode_maintenance_reasons, fetch_metars
from asos_tools.stations import AOMC_STATIONS

__all__ = ["generate_incident_docx"]


# ----------------------------------------------------------------------------
# Style constants
# ----------------------------------------------------------------------------

NOAA_BLUE = RGBColor(0, 51, 102)
HEADER_GRAY = RGBColor(80, 80, 80)
TEXT_BLACK = RGBColor(30, 30, 30)
TABLE_HEADER_BG = "003366"
TABLE_ALT_ROW = "F0F4F8"
RED_FLAG = RGBColor(180, 30, 30)
GREEN_CLEAN = RGBColor(20, 120, 60)
AMBER = RGBColor(180, 120, 0)
HARD_FAIL_BG = "FEE2E2"

FONT_BODY = "Calibri"
FONT_MONO = "Consolas"


# ----------------------------------------------------------------------------
# Small helpers
# ----------------------------------------------------------------------------

def _shade(cell, hex_color: str) -> None:
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    )


def _narrow_margins(section) -> None:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)


def _add_rule(doc) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p._p.get_or_add_pPr().append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="003366"/>'
        '</w:pBdr>'
    ))


def _header_row(table, texts: list[str]) -> None:
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = FONT_BODY
        run.font.color.rgb = RGBColor(255, 255, 255)
        _shade(cell, TABLE_HEADER_BG)


def _data_row(table, values, *, row_idx: int = 0,
              highlight_col: int | None = None,
              highlight_color: RGBColor | None = None) -> None:
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(val))
        run.font.size = Pt(8.5)
        run.font.name = FONT_BODY
        run.font.color.rgb = TEXT_BLACK
        if highlight_col is not None and i == highlight_col and highlight_color:
            run.font.color.rgb = highlight_color
            run.bold = True
        if row_idx % 2 == 0:
            _shade(cell, TABLE_ALT_ROW)


def _meta_line(p, label: str, value: str) -> None:
    run_l = p.add_run(f"{label}  ")
    run_l.bold = True
    run_l.font.size = Pt(9)
    run_l.font.name = FONT_BODY
    run_l.font.color.rgb = NOAA_BLUE
    run_v = p.add_run(f"{value}\n")
    run_v.font.size = Pt(9)
    run_v.font.name = FONT_BODY
    run_v.font.color.rgb = TEXT_BLACK


def _metar_callout(doc, label: str, metar: str, *, severity: str = "info") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    cell.text = ""

    if severity == "hard":
        _shade(cell, HARD_FAIL_BG)
        accent = RED_FLAG
    elif severity == "warn":
        _shade(cell, "FEF3C7")
        accent = AMBER
    else:
        _shade(cell, "F1F5F9")
        accent = NOAA_BLUE

    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    run = p1.add_run(f"{label}\n")
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = FONT_BODY
    run.font.color.rgb = accent

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    run = p2.add_run(metar)
    run.font.size = Pt(9)
    run.font.name = FONT_MONO
    run.font.color.rgb = TEXT_BLACK


# ----------------------------------------------------------------------------
# Per-station analysis
# ----------------------------------------------------------------------------

def _investigate(stn: str, hours: float, end: datetime) -> dict:
    start = end - timedelta(hours=hours)
    df = fetch_metars(stn, start, end)
    if df.empty:
        return {"station": stn, "empty": True}

    df = df.sort_values("valid").reset_index(drop=True)
    total = len(df)
    flagged = df[df["has_maintenance"]].copy()
    n_flag = len(flagged)

    incidents = []
    in_inc = False
    cur_rows: list = []
    cur_start = None
    for _, row in df.iterrows():
        if row["has_maintenance"]:
            if not in_inc:
                in_inc = True
                cur_start = row["valid"]
                cur_rows = []
            cur_rows.append(row)
        else:
            if in_inc:
                incidents.append({
                    "start": cur_start,
                    "end": cur_rows[-1]["valid"],
                    "rows": cur_rows,
                    "duration_hr": (cur_rows[-1]["valid"] - cur_start)
                    .total_seconds() / 3600,
                })
                in_inc = False
    if in_inc:
        incidents.append({
            "start": cur_start,
            "end": cur_rows[-1]["valid"],
            "rows": cur_rows,
            "duration_hr": (cur_rows[-1]["valid"] - cur_start).total_seconds() / 3600,
            "ongoing": True,
        })

    sensor_counts: dict[str, int] = {}
    hard_failures = []
    for _, row in flagged.iterrows():
        for r in decode_maintenance_reasons(row["metar"]):
            sensor_counts[r["sensor"]] = sensor_counts.get(r["sensor"], 0) + 1
            if r["sensor"] != "Internal check":
                hard_failures.append(r)

    return {
        "station": stn,
        "empty": False,
        "total": total,
        "flagged": n_flag,
        "rate": n_flag / total * 100 if total else 0,
        "incidents": incidents,
        "sensor_counts": sensor_counts,
        "hard_failures": hard_failures,
    }


# ----------------------------------------------------------------------------
# Public API
# ----------------------------------------------------------------------------

def generate_incident_docx(
    stations: Iterable[str],
    *,
    hours: float = 168,
    end: datetime | None = None,
    out: Union[str, Path, io.BytesIO, None] = None,
) -> bytes:
    """Generate a NOAA-styled DOCX incident investigation report.

    Parameters
    ----------
    stations
        ICAO station IDs (e.g. ``["KGSO", "KFAY"]``).
    hours
        Lookback window in hours (default 168 = 7 days).
    end
        End of the investigation window. Defaults to "now" (UTC).
    out
        Where to write the document. Accepts a filesystem path
        (``str`` or ``Path``), an ``io.BytesIO``, or ``None``. When
        ``None``, the document is buffered in memory and the raw bytes
        are returned.

    Returns
    -------
    bytes
        The DOCX bytes. If ``out`` is a path the file is also written there.
    """
    stations = [s.upper() for s in stations]
    end = end or datetime.now(timezone.utc)
    start = end - timedelta(hours=hours)

    investigations = [_investigate(s, hours, end) for s in stations]

    total_metars = sum(i.get("total", 0) for i in investigations)
    total_flagged = sum(i.get("flagged", 0) for i in investigations)
    total_incidents = sum(len(i.get("incidents", [])) for i in investigations)
    total_hard_fails = sum(len(i.get("hard_failures", [])) for i in investigations)

    doc = Document()
    _narrow_margins(doc.sections[0])

    # --- Title block ---
    title = doc.add_heading("ASOS Maintenance-Flag Incident Investigation", level=1)
    title.runs[0].font.color.rgb = NOAA_BLUE

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(2)
    r = sub.add_run("Automated Surface Observing System  |  "
                    "$ Maintenance Indicator Trace Analysis")
    r.font.size = Pt(10)
    r.font.color.rgb = HEADER_GRAY
    r.font.name = FONT_BODY

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(6)
    _meta_line(meta, "Report Generated:", end.strftime("%Y-%m-%dT%H:%M:%SZ"))
    _meta_line(meta, "Investigation Window:",
               f"{start.strftime('%Y-%m-%dT%H:%MZ')} -> "
               f"{end.strftime('%Y-%m-%dT%H:%MZ')}  ({hours:.0f} hours)")
    _meta_line(meta, "Stations Investigated:",
               f"{', '.join(stations)}  ({len(stations)} sites)")
    _meta_line(meta, "Data Source:",
               "NOAA / NCEI ASOS METAR archive via Iowa Environmental Mesonet (IEM)")
    _meta_line(meta, "Decoder:",
               "asos_tools.metars (RVRNO, PWINO, PNO, FZRANO, TSNO, "
               "VISNO, CHINO, SLPNO)")
    _add_rule(doc)

    # --- Executive summary ---
    doc.add_heading("Executive Summary", level=2).runs[0].font.color.rgb = NOAA_BLUE
    sumt = doc.add_table(rows=1, cols=5)
    sumt.alignment = WD_TABLE_ALIGNMENT.CENTER
    sumt.style = "Table Grid"
    _header_row(sumt, ["Stations", "METARs Reviewed", "Flagged ($)",
                       "Distinct Incidents", "Hard Sensor Failures"])
    _data_row(sumt, [len(stations), f"{total_metars:,}", f"{total_flagged:,}",
                     total_incidents, total_hard_fails], row_idx=0)
    data_row = sumt.rows[1]
    colors = [NOAA_BLUE, TEXT_BLACK, RED_FLAG, RED_FLAG,
              RED_FLAG if total_hard_fails else GREEN_CLEAN]
    for cell, color in zip(data_row.cells, colors):
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = color
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    p = doc.add_paragraph()
    if total_hard_fails > 0:
        run = p.add_run(
            f"FINDING:  {total_hard_fails} hard sensor failure(s) and "
            f"{total_incidents - total_hard_fails} tolerance-check incident(s) "
            f"detected across {len(stations)} station(s). Each is a distinct "
            "$ event that downstream ticketing systems should have processed.")
        run.font.color.rgb = RED_FLAG
        run.bold = True
    elif total_incidents > 0:
        run = p.add_run(
            f"FINDING:  {total_incidents} tolerance-check incident(s) detected, "
            "no explicit hard-sensor failures. Each should still be processed by "
            "downstream ticketing.")
        run.font.color.rgb = AMBER
        run.bold = True
    else:
        run = p.add_run(
            "FINDING:  No $ maintenance flags detected in this window. "
            "Downstream silence on these sites is correct.")
        run.font.color.rgb = GREEN_CLEAN
        run.bold = True
    run.font.size = Pt(10)
    _add_rule(doc)

    # --- Per-station ---
    for inv in investigations:
        stn = inv["station"]
        meta_record = next((s for s in AOMC_STATIONS if s.get("id") == stn), None)
        name = meta_record.get("name", "") if meta_record else ""
        heading = doc.add_heading(
            f"{stn}  -  {name}".strip(" -"), level=2)
        heading.runs[0].font.color.rgb = NOAA_BLUE

        if inv.get("empty"):
            p = doc.add_paragraph()
            r = p.add_run(f"No METARs returned for {stn} in this window.")
            r.font.color.rgb = HEADER_GRAY
            r.font.size = Pt(9.5)
            _add_rule(doc)
            continue

        chip = doc.add_table(rows=1, cols=4)
        chip.style = "Table Grid"
        _header_row(chip, ["METARs Received", "Flagged ($)", "Flag Rate",
                           "Distinct Incidents"])
        _data_row(chip, [f"{inv['total']:,}", f"{inv['flagged']:,}",
                         f"{inv['rate']:.1f}%", len(inv["incidents"])],
                  row_idx=0)
        rate_cell = chip.rows[1].cells[2]
        for run in rate_cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = (RED_FLAG if inv["rate"] >= 50
                                  else AMBER if inv["rate"] >= 10
                                  else GREEN_CLEAN)
        inc_cell = chip.rows[1].cells[3]
        for run in inc_cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RED_FLAG if inv["incidents"] else GREEN_CLEAN

        doc.add_paragraph("")

        if inv["incidents"]:
            p = doc.add_paragraph()
            r = p.add_run("Incident Timeline")
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = NOAA_BLUE

            for idx, inc in enumerate(inv["incidents"], start=1):
                inc_sensors = set()
                for row in inc["rows"]:
                    for code in decode_maintenance_reasons(row["metar"]):
                        inc_sensors.add(code["sensor"])
                sensors_str = ", ".join(sorted(inc_sensors))

                p_inc = doc.add_paragraph()
                p_inc.paragraph_format.space_before = Pt(4)
                p_inc.paragraph_format.space_after = Pt(2)
                r_t = p_inc.add_run(f"Incident #{idx}  ")
                r_t.bold = True
                r_t.font.size = Pt(9.5)
                r_t.font.color.rgb = (RED_FLAG
                                      if len(inc_sensors) > 1 or "Internal check" not in inc_sensors
                                      else AMBER)
                dur_h = inc["duration_hr"]
                dur_str = ("1 METAR cycle" if dur_h < 1
                           else f"{dur_h:.1f} hours" if dur_h < 24
                           else f"{dur_h/24:.1f} days")
                r_d = p_inc.add_run(
                    f"{inc['start'].strftime('%Y-%m-%d %H:%MZ')}  to  "
                    f"{inc['end'].strftime('%Y-%m-%d %H:%MZ')}  "
                    f"({dur_str})  -  {sensors_str}"
                    + ("  [ONGOING]" if inc.get("ongoing") else ""))
                r_d.font.size = Pt(9.5)
                r_d.font.color.rgb = TEXT_BLACK

                severity = ("hard"
                            if inc_sensors and "Internal check" not in inc_sensors
                            else "warn")
                _metar_callout(doc,
                               f"First $ METAR of incident #{idx}",
                               inc["rows"][0]["metar"],
                               severity=severity)
                doc.add_paragraph("")

        if inv["sensor_counts"]:
            p = doc.add_paragraph()
            r = p.add_run("Sensor Code Breakdown")
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = NOAA_BLUE

            st_tbl = doc.add_table(rows=1, cols=3)
            st_tbl.style = "Table Grid"
            _header_row(st_tbl, ["Sensor / Subsystem", "Affected METARs", "Severity"])
            for idx, (sensor, count) in enumerate(
                sorted(inv["sensor_counts"].items(), key=lambda x: -x[1])
            ):
                if sensor == "Internal check":
                    sev, color = "Tolerance Check", AMBER
                else:
                    sev, color = "HARD SENSOR FAILURE", RED_FLAG
                _data_row(st_tbl, [sensor, count, sev],
                          row_idx=idx, highlight_col=2, highlight_color=color)
            doc.add_paragraph("")

        _add_rule(doc)

    # --- Root Cause Analysis ---
    doc.add_heading("Root Cause Analysis: Missing Downstream Tickets",
                    level=2).runs[0].font.color.rgb = NOAA_BLUE
    for label, text in [
        ("Bouncy $ flag behavior. ",
         "Routine METARs (filed near HH:51Z) carry the $ flag; intervening "
         "SPECIs often do not. Ingesters subscribed only to SPECIs miss the "
         "flag."),
        ("Single-cycle flag suppression. ",
         "Incidents that clear within one METAR cycle are silently dropped "
         "by ingesters that require N>=2 consecutive flagged reports."),
        ("Hard-failure code stripping. ",
         "If the parser strips the literal $ before flag detection, no "
         "incident is raised regardless of RVRNO / PWINO / TSNO presence."),
        ("Network-wide event de-duplication. ",
         "Multiple stations flagging in the same minute may be deduped as "
         "one network event, dropping per-site tickets."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r_l = p.add_run(label)
        r_l.bold = True
        r_l.font.size = Pt(9)
        r_l.font.color.rgb = NOAA_BLUE
        r_v = p.add_run(text)
        r_v.font.size = Pt(9)
        r_v.font.color.rgb = TEXT_BLACK
    _add_rule(doc)

    # --- Recommendations ---
    doc.add_heading("Recommendations", level=2).runs[0].font.color.rgb = NOAA_BLUE
    recs = [
        "Verify the ingester reads routine METARs (TYPE=METAR) as well as SPECIs.",
        "Audit the parser for handling of trailing $ - confirm it is preserved "
        "through the full pipeline.",
        "Use first-flag-in-window dedup (one ticket per station per 4-hour "
        "window) rather than per-METAR ticketing.",
        "Parse explicit sensor codes (RVRNO, PWINO, PNO, FZRANO, TSNO, VISNO, "
        "CHINO, SLPNO) and escalate hard-sensor failures above tolerance checks.",
        "Cross-reference this report against downstream tickets for the same "
        "window. Confirm each incident listed above created a corresponding "
        "ticket.",
    ]
    for i, rec in enumerate(recs, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r_n = p.add_run(f"{i}.  ")
        r_n.bold = True
        r_n.font.size = Pt(9)
        r_n.font.color.rgb = NOAA_BLUE
        r_t = p.add_run(rec)
        r_t.font.size = Pt(9)
        r_t.font.color.rgb = TEXT_BLACK
    _add_rule(doc)

    # --- Methodology ---
    doc.add_heading("Methodology", level=2).runs[0].font.color.rgb = NOAA_BLUE
    for label, text in [
        ("METAR Source: ",
         "Iowa Environmental Mesonet ASOS service "
         "(mesonet.agron.iastate.edu/cgi-bin/request/asos.py), report_type=3."),
        ("Flag Detection: ",
         "A METAR is flagged when the trimmed report ends with the $ "
         "character (after stripping any trailing =)."),
        ("Sensor Code Decoding: ",
         "Remarks scanned for RVRNO, PWINO, PNO, FZRANO, TSNO, VISNO [loc], "
         "CHINO [loc], SLPNO per the ASOS User's Guide and FMH-1 Table 8-5. "
         "Flagged METARs without explicit codes are classified 'Internal check'."),
        ("Incident Definition: ",
         "A contiguous run of flagged METARs at one station, terminated by a "
         "clean METAR or the end of the window."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r_l = p.add_run(label)
        r_l.bold = True
        r_l.font.size = Pt(8.5)
        r_l.font.color.rgb = NOAA_BLUE
        r_t = p.add_run(text)
        r_t.font.size = Pt(8.5)
        r_t.font.color.rgb = TEXT_BLACK
    _add_rule(doc)

    # --- Footer ---
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(8)
    r = footer.add_run(
        "O.W.L. - Observation Watch Log  |  Incident Investigation Module  |  "
        "Data sourced from NOAA/NCEI via the Iowa Environmental Mesonet.")
    r.font.size = Pt(7.5)
    r.font.color.rgb = HEADER_GRAY
    r.italic = True

    # --- Serialize ---
    buffer = io.BytesIO()
    doc.save(buffer)
    data = buffer.getvalue()

    if isinstance(out, (str, Path)):
        Path(out).parent.mkdir(parents=True, exist_ok=True)
        Path(out).write_bytes(data)
    elif isinstance(out, io.BytesIO):
        out.write(data)

    return data
